from __future__ import annotations

import os
import random
import re
from datetime import date, datetime
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.table import _Cell


COLUMN_MAPPING = {
    "code": ["代码", "股票代码", "证券代码"],
    "name": ["名称", "股票名称", "证券简称"],
    "inquiry_date": ["询价日", "询价日期", "日期"],
}

PERSONS = ["韩震泓", "余晓舰", "王博", "罗钰瑶"]
INVALID_FILENAME_CHARS = r'<>:"/\|?*'


def normalize_header(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"[\s:：]+", "", str(value).strip())


def clean_text(value: Any) -> str:
    if value is None or pd.isna(value):
        return ""
    text = str(value).strip()
    if text in {"nan", "None"}:
        return ""
    return text


def format_stock_code(value: Any) -> str:
    if value is None or pd.isna(value):
        return ""

    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return str(int(value)).zfill(6)

    text = clean_text(value)
    if not text:
        return ""

    if re.fullmatch(r"\d+(\.0+)?", text):
        return str(int(float(text))).zfill(6)

    six_digit_match = re.search(r"\d{6}", text)
    if six_digit_match:
        return six_digit_match.group(0)

    return text.zfill(6) if text.isdigit() else text


def parse_inquiry_date(value: Any) -> tuple[str, str]:
    if value is None or pd.isna(value):
        return "", ""

    if isinstance(value, pd.Timestamp):
        value = value.to_pydatetime()

    if isinstance(value, datetime):
        return value.strftime("%Y年%m月%d日"), value.strftime("%Y%m%d")

    if isinstance(value, date):
        return value.strftime("%Y年%m月%d日"), value.strftime("%Y%m%d")

    text = clean_text(value)
    if not text:
        return "", ""

    if re.fullmatch(r"\d{8}", text):
        try:
            parsed = datetime.strptime(text, "%Y%m%d")
            return parsed.strftime("%Y年%m月%d日"), parsed.strftime("%Y%m%d")
        except ValueError:
            return text, text

    for fmt in ("%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y年%m月%d日"):
        try:
            parsed = datetime.strptime(text, fmt)
            return parsed.strftime("%Y年%m月%d日"), parsed.strftime("%Y%m%d")
        except ValueError:
            continue

    return text, re.sub(r"\D+", "", text) or safe_filename_part(text)


def safe_filename_part(value: Any) -> str:
    text = clean_text(value)
    cleaned = "".join(ch for ch in text if ch not in INVALID_FILENAME_CHARS)
    return cleaned.strip(" .") or "未命名"


def unique_path(path: Path) -> Path:
    if not path.exists():
        return path

    stem = path.stem
    suffix = path.suffix
    parent = path.parent
    index = 2
    while True:
        candidate = parent / f"{stem}_{index}{suffix}"
        if not candidate.exists():
            return candidate
        index += 1


def set_cell_border(cell: _Cell) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()

    for side in ["left", "right", "top", "bottom"]:
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")
        tc_pr.append(border)


def apply_run_font(run) -> None:
    run.font.name = "仿宋"
    run.font.size = Pt(15)
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")


def find_all_excel_files(root_dir: str | os.PathLike[str]) -> list[Path]:
    excel_files: list[Path] = []
    for root, _, files in os.walk(root_dir):
        for file_name in files:
            if file_name.startswith("~$"):
                continue
            if file_name.lower().endswith((".xlsx", ".xls")):
                excel_files.append(Path(root) / file_name)

    if not excel_files:
        raise FileNotFoundError(f"在目录 '{root_dir}' 及其子目录中未找到 Excel 文件")

    return excel_files


def find_columns(columns: list[Any]) -> dict[str, Any]:
    normalized = {normalize_header(column): column for column in columns}
    found: dict[str, Any] = {}

    for target_key, possible_names in COLUMN_MAPPING.items():
        for name in possible_names:
            column = normalized.get(normalize_header(name))
            if column is not None:
                found[target_key] = column
                break

    return found


def load_excel_with_header(excel_file: Path) -> tuple[pd.DataFrame, dict[str, Any]]:
    errors: list[str] = []
    for header_row in range(0, 7):
        try:
            df = pd.read_excel(excel_file, header=header_row, dtype=object)
        except Exception as exc:
            errors.append(str(exc))
            continue

        found_columns = find_columns(list(df.columns))
        if all(key in found_columns for key in COLUMN_MAPPING):
            return df, found_columns

    detail = f"；读取错误：{errors[0]}" if errors else ""
    raise ValueError(f"无法找到必要列：代码、名称、询价日。文件路径：{excel_file}{detail}")


def extract_records(excel_file: str | os.PathLike[str]) -> list[dict[str, str]]:
    path = Path(excel_file)
    df, columns = load_excel_with_header(path)
    records: list[dict[str, str]] = []

    for _, row in df.iterrows():
        code = format_stock_code(row.get(columns["code"]))
        name = clean_text(row.get(columns["name"]))
        date_text, date_file = parse_inquiry_date(row.get(columns["inquiry_date"]))

        if not code or not name or not date_text:
            continue

        records.append(
            {
                "code": code,
                "name": name,
                "date_text": date_text,
                "date_file": date_file,
            }
        )

    if not records:
        raise ValueError(f"Excel 文件中没有有效数据。文件路径：{excel_file}")

    return records


def create_word_document(record: dict[str, str], output_dir: str | os.PathLike[str]) -> Path:
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "仿宋"
    style.font.element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋")
    style.font.size = Pt(15)
    style.paragraph_format.line_spacing = 1.5

    title = doc.add_paragraph("新股询价现场通讯工具上交登记表")
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    apply_run_font(title.runs[0])

    doc.add_paragraph(f"股票名称：【{record['name']}】")
    doc.add_paragraph(f"股票代码：【{record['code']}】")
    doc.add_paragraph(f"询价日期：{record['date_text']}")
    doc.add_paragraph("询价关键时间窗口：09：30 - 15：00")
    doc.add_paragraph("保管地点：【合规部】")
    doc.add_paragraph()

    table = doc.add_table(rows=5, cols=5)
    headers = ["人员", "上交时间", "通讯工具类型", "上交确认", "领取时间"]
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                apply_run_font(run)
        set_cell_border(cell)

    for row_index, person in enumerate(PERSONS, start=1):
        hand_in_time = f"09：{random.randint(0, 29):02d}"
        pick_up_time = f"15：{random.randint(1, 30):02d}"
        values = [person, hand_in_time, "手机", "√", pick_up_time]

        for col_index, value in enumerate(values):
            cell = table.cell(row_index, col_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    apply_run_font(run)
            set_cell_border(cell)

    filename = (
        f"{safe_filename_part(record['code'])}_"
        f"{safe_filename_part(record['name'])}_"
        f"{safe_filename_part(record['date_file'])}_通讯工具上交登记表.docx"
    )
    doc_path = unique_path(output_path / filename)
    doc.save(doc_path)
    return doc_path


def process_excel_files(input_dir: str | os.PathLike[str], output_dir: str | os.PathLike[str]) -> dict[str, Any]:
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    excel_files = find_all_excel_files(input_dir)
    stats: dict[str, Any] = {
        "excel_total": len(excel_files),
        "excel_processed": 0,
        "records_total": 0,
        "generated_total": 0,
        "generated_files": [],
        "errors": [],
    }

    for excel_file in excel_files:
        try:
            records = extract_records(excel_file)
            stats["excel_processed"] += 1
            stats["records_total"] += len(records)

            for record in records:
                doc_path = create_word_document(record, output_path)
                stats["generated_total"] += 1
                stats["generated_files"].append(str(doc_path))
        except Exception as exc:
            stats["errors"].append(f"{excel_file.name}: {exc}")

    if stats["generated_total"] == 0:
        details = "; ".join(stats["errors"][:5]) if stats["errors"] else "未生成任何 Word 文档"
        raise RuntimeError(details)

    return stats


if __name__ == "__main__":
    script_dir = Path(__file__).resolve().parent
    result = process_excel_files(script_dir, script_dir)
    print(f"成功生成 {result['generated_total']} 个文件")

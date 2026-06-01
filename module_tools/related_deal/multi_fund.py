import os
import re
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
import xlrd
from datetime import datetime, timedelta
from docx.oxml.ns import qn

# 客户名称和客户代码对应关系
client_code_map = {
    '睿量智行1号私募证券投资基金': 'SXE629',
    '睿量智行1号1期私募证券投资基金': 'SACK10',
    '睿量智行3号私募证券投资基金': 'SZN727',
    '睿量电子1号私募证券投资基金': 'SXZ194',
    '睿量电子3号私募证券投资基金': 'SZT965',
    '睿量电子11号私募证券投资基金': 'SZT155',
    '睿量原子1号私募证券投资基金': 'SXG468',
    '睿量电子18号私募证券投资基金': 'SAFH20',
    '睿量电子7号私募证券投资基金': 'SZV042',
    '睿量智享1号私募证券投资基金': 'SXE782',
    '睿量红星2号私募证券投资基金': 'SZT967',
    '睿量智享2号私募证券投资基金': 'SAEF48',
    '睿量光子1号私募证券投资基金': 'SZA093',
    '睿量光子1号1期私募证券投资基金': 'SACQ29',
    '睿量中微子1号私募证券投资基金': 'SXY950',
    '睿量分子1号私募证券投资基金': 'SXQ540',
    '睿量中子1号私募证券投资基金': 'SXQ538',
    '睿量中子3号私募证券投资基金': 'SB4362',
    '睿量红星1号私募证券投资基金': 'SXY949',
    '睿量中子2号私募证券投资基金': 'SZR863',
    '睿量电子2号私募证券投资基金': 'SZT964',
    '睿量电子6号私募证券投资基金': 'SZV040',
    '睿量智行2号私募证券投资基金': 'SZN726',
    '睿量量子9号私募证券投资基金': 'SZD411',
    '睿量量子春晓1号私募证券投资基金': 'SZT700',
    '睿量中微子2号私募证券投资基金': 'SAHZ35',
    '睿量电子5号私募证券投资基金': 'SZV038',
    '睿量量子3号私募证券投资基金': 'SZB168',
    '睿量电子9号私募证券投资基金': 'SZW081',
    '睿量量子1号私募证券投资基金': 'SZA134',
    '睿量量子11号私募证券投资基金': 'SZN739',
    '睿量电子12号私募证券投资基金': 'SB8147',
    '睿量电子13号私募证券投资基金': 'SB8309',
    '睿量福兴2号私募证券投资基金': 'SB8370',
    '睿量光子春晓2号私募证券投资基金': 'SABU17',
    '睿量电子15号私募证券投资基金': 'SACZ40',
    '睿量智行广睿1号私募证券投资基金': 'SADE31',
    '睿量智行7号私募证券投资基金': 'SAEP50',
    '睿量智行7号私募证券投资基金1期': 'SAEP51',
    '睿量智行7号私募证券投资基金2期': 'SAEP52',
    '睿量光子6号私募证券投资基金': 'SADX94',
    '睿量光子6号一期私募证券投资基金': 'SADX95',
    '睿量光子6号二期私募证券投资基金': 'SADX97',
    '睿量红光1号私募证券投资基金': 'SAES61',
    '省心享睿量兴泰奋进1期私募证券投资基金': 'SAGA05',
    '睿量量子12号私募证券投资基金': 'SAES62',
    '睿量量子13号私募证券投资基金': 'SAES63',
    '睿量臻选1000指数增强1号私募证券投资基金': 'SAES60',
    '睿量信淮智行9号私募证券投资基金': 'SAFD11',
    '睿量信淮智行9号1期私募证券投资基金': 'SAFD12',
    '睿量信淮智行9号2期私募证券投资基金': 'SAFF49',
    '省心享睿量兴泰锐进1号私募证券投资基金': 'SAGC89',
    '睿量电子21号私募证券投资基金': 'SAGA28',
    '睿量质子1号私募证券投资基金': 'SXT232',
    '睿量量子聚利1号私募证券投资基金': 'SAJE31',
    '睿量电子17号私募证券投资基金': 'SAJQ36',
    '睿量量子15号私募证券投资基金': 'SAKA78',
    '睿量量子10号私募证券投资基金': 'SZN736',
    '睿量电子22号私募证券投资基金': 'SAKT70',
    '睿量原子10号私募证券投资基金': 'SALP86',
    '睿量原子5号私募证券投资基金': 'SARF66',
    '睿量原子2号私募证券投资基金': 'SB2710',
    '睿量原子2号一期私募证券投资基金': 'SAJM39',
    '睿量原子2号二期私募证券投资基金': 'SAJM40',
    '睿量陆享原子6号私募证券投资基金': 'SARM66',
    '睿量中远1号私募证券投资基金': 'SARY23',
    '睿量原子8号私募证券投资基金': 'SASP78',
    '睿量量子聚利1号1期私募证券投资基金': 'SAPM39',
    '睿量量子春晓3号私募证券投资基金': 'SASR13',
    '睿量陆享原子16号私募证券投资基金': 'SASZ91',
    '睿量原子7号私募证券投资基金': 'SATB02',
    '睿量原子7号私募证券投资基金1期': 'SATA99',
    '睿量原子7号私募证券投资基金2期': 'SATB01',
    '睿量信淮原子11号私募证券投资基金': 'SATJ62',
    '睿量信淮原子9号私募证券投资基金': 'SATJ30',
    '睿量信淮原子9号1期私募证券投资基金': 'SATJ26',
    '睿量HT鹏泰1号私募证券投资基金': 'SAUA14',
    '睿量量子-江海远山51期私募证券投资基金': 'SAUH59',
    '睿量信选原子3号私募证券投资基金': 'SATL24',
    '睿量电子20号私募证券投资基金': 'SAGA23'
}

def find_all_excel_files(root_dir='.'):
    """递归查找所有文件夹中的Excel文件"""
    excel_files = []
    
    for root, _, files in os.walk(root_dir):
        for file in files:
            if file.lower().endswith(('.xlsx', '.xls')):
                full_path = os.path.join(root, file)
                excel_files.append(full_path)
    
    if not excel_files:
        raise FileNotFoundError(f"在目录 '{root_dir}' 及其子目录中未找到Excel文件")
    
    return excel_files

def extract_info(excel_file):
    """从Excel文件中提取信息"""
    all_records = []
    
    try:
        # 根据文件扩展名选择读取方法
        file_ext = os.path.splitext(excel_file)[1].lower()
        
        if file_ext in ['.xlsx', '.xlsm']:
            # 使用openpyxl读取xlsx文件
            wb = openpyxl.load_workbook(excel_file, data_only=True)
            
            for sheet_name in wb.sheetnames:
                print(f"处理工作表: {sheet_name}")
                ws = wb[sheet_name]
                
                # 尝试不同的表头行数
                for header_row in range(1, 11):  # 尝试前10行作为表头
                    try:
                        # 查找列名
                        found_columns = {}
                        column_mapping = {
                            'client_name': ['客户名称', '投资者名称'],
                            'apply_date': ['申请日期'],
                            'apply_statue': ['确认状态', '确认情况']
                        }
                        
                        # 遍历表头行，查找列名
                        for col in range(1, ws.max_column + 1):
                            cell_value = ws.cell(row=header_row, column=col).value
                            if cell_value:
                                header = str(cell_value).strip()
                                for key, possible_names in column_mapping.items():
                                    if key not in found_columns:
                                        # 只使用精确匹配，避免将日期列错误识别为状态列
                                        if header in possible_names:
                                            found_columns[key] = col - 1  # 转换为0-based索引
                        
                        required_columns = ['client_name', 'apply_date']
                        if all(key in found_columns for key in required_columns):
                            print(f"找到所有列名: {found_columns}")

                            # 处理数据行
                            for row in range(header_row + 1, ws.max_row + 1):
                                # 处理客户名称
                                client_cell = ws.cell(row=row, column=found_columns['client_name'] + 1)
                                client_name = str(client_cell.value).strip() if client_cell.value else ''

                                if not client_name or client_name in ['nan', 'None']:
                                    continue

                                # 处理申请日期
                                date_cell = ws.cell(row=row, column=found_columns['apply_date'] + 1)
                                date_val = date_cell.value
                                apply_date = None

                                if date_val:
                                    try:
                                        # 处理Excel日期对象
                                        if isinstance(date_val, datetime):
                                            apply_date = date_val.strftime('%Y年%m月%d日')
                                        # 处理Excel序列号
                                        elif isinstance(date_val, (int, float)):
                                            # Excel日期序列号从1900年1月1日开始
                                            base_date = datetime(1899, 12, 30)
                                            converted_date = base_date + timedelta(days=date_val)
                                            # 检查日期是否合理（大于2000年）
                                            if converted_date.year > 2000:
                                                apply_date = converted_date.strftime('%Y年%m月%d日')
                                        # 处理字符串格式
                                        else:
                                            date_str = str(date_val).strip().replace(' ', '').replace(',', '').replace('¥', '').replace('￥', '')
                                            # 尝试yyyymmdd格式
                                            if len(date_str) == 8 and date_str.isdigit():
                                                apply_date = f"{date_str[:4]}年{date_str[4:6]}月{date_str[6:]}日"
                                            else:
                                                # 尝试其他日期格式
                                                date_formats = ['%Y-%m-%d', '%Y/%m/%d', '%d/%m/%Y', '%d-%m-%Y', '%Y年%m月%d日', '%Y.%m.%d']
                                                for fmt in date_formats:
                                                    try:
                                                        apply_date = datetime.strptime(date_str, fmt).strftime('%Y年%m月%d日')
                                                        break
                                                    except Exception:
                                                        continue
                                    except Exception:
                                        apply_date = str(date_val).strip()

                                if not apply_date:
                                    continue

                                # 处理确认状态（可选列，缺失时默认为空）
                                apply_statue = ''
                                if 'apply_statue' in found_columns:
                                    status_cell = ws.cell(row=row, column=found_columns['apply_statue'] + 1)
                                    apply_statue = str(status_cell.value).strip() if status_cell.value else ''

                                record = {
                                    'client_name': client_name,
                                    'apply_date': apply_date,
                                    'apply_statue': apply_statue
                                }

                                all_records.append(record)

                            # 找到有效数据后，跳出循环
                            break
                        else:
                            print(f"未找到所有列名，找到的列名: {found_columns}")
                    except Exception as e:
                        print(f"尝试表头行 {header_row} 失败: {e}")
                        import traceback
                        traceback.print_exc()
                        continue
        
        elif file_ext == '.xls':
            # 使用xlrd读取xls文件
            wb = xlrd.open_workbook(excel_file)
            
            for sheet_name in wb.sheet_names():
                print(f"处理工作表: {sheet_name}")
                ws = wb.sheet_by_name(sheet_name)
                
                # 尝试不同的表头行数
                for header_row in range(0, min(10, ws.nrows)):  # 尝试前10行作为表头
                    try:
                        # 查找列名
                        found_columns = {}
                        column_mapping = {
                            'client_name': ['客户名称', '投资者名称', '客户', '基金名称', '产品名称'],
                            'apply_date': ['申请日期', '日期', '交易日期', '确认日期'],
                            'apply_statue': ['确认状态', '状态', '交易状态', '确认']
                        }
                        
                        # 遍历表头行，查找列名
                        for col in range(0, ws.ncols):
                            cell_value = ws.cell_value(header_row, col)
                            if cell_value:
                                header = str(cell_value).strip()
                                for key, possible_names in column_mapping.items():
                                    if key not in found_columns:
                                        # 只使用精确匹配，避免将日期列错误识别为状态列
                                        if header in possible_names:
                                            found_columns[key] = col
                        
                        required_columns = ['client_name', 'apply_date']
                        if all(key in found_columns for key in required_columns):
                            print(f"找到所有列名: {found_columns}")

                            # 处理数据行
                            for row in range(header_row + 1, ws.nrows):
                                # 处理客户名称
                                client_value = ws.cell_value(row, found_columns['client_name'])
                                client_name = str(client_value).strip() if client_value else ''

                                if not client_name or client_name in ['nan', 'None']:
                                    continue

                                # 处理申请日期
                                date_value = ws.cell_value(row, found_columns['apply_date'])
                                apply_date = None

                                if date_value:
                                    try:
                                        # 处理Excel日期序列号
                                        if isinstance(date_value, (int, float)):
                                            # xlrd使用不同的日期处理
                                            if date_value > 0:
                                                # 尝试转换为日期
                                                try:
                                                    date_tuple = xlrd.xldate_as_tuple(date_value, wb.datemode)
                                                    if date_tuple[0] > 2000:  # 确保年份合理
                                                        apply_date = f"{date_tuple[0]}年{date_tuple[1]:02d}月{date_tuple[2]:02d}日"
                                                except Exception:
                                                    pass
                                        # 处理字符串格式
                                        else:
                                            date_str = str(date_value).strip().replace(' ', '').replace(',', '').replace('¥', '').replace('￥', '')
                                            # 尝试yyyymmdd格式
                                            if len(date_str) == 8 and date_str.isdigit():
                                                apply_date = f"{date_str[:4]}年{date_str[4:6]}月{date_str[6:]}日"
                                            else:
                                                # 尝试其他日期格式
                                                date_formats = ['%Y-%m-%d', '%Y/%m/%d', '%d/%m/%Y', '%d-%m-%Y', '%Y年%m月%d日', '%Y.%m.%d']
                                                for fmt in date_formats:
                                                    try:
                                                        apply_date = datetime.strptime(date_str, fmt).strftime('%Y年%m月%d日')
                                                        break
                                                    except Exception:
                                                        continue
                                    except Exception:
                                        apply_date = str(date_value).strip()

                                if not apply_date:
                                    continue

                                # 处理确认状态（可选列，缺失时默认为空）
                                apply_statue = ''
                                if 'apply_statue' in found_columns:
                                    status_value = ws.cell_value(row, found_columns['apply_statue'])
                                    apply_statue = str(status_value).strip() if status_value else ''

                                record = {
                                    'client_name': client_name,
                                    'apply_date': apply_date,
                                    'apply_statue': apply_statue
                                }

                                all_records.append(record)

                            # 找到有效数据后，跳出循环
                            break
                        else:
                            print(f"未找到所有列名，找到的列名: {found_columns}")
                    except Exception as e:
                        print(f"尝试表头行 {header_row} 失败: {e}")
                        import traceback
                        traceback.print_exc()
                        continue
        else:
            print(f"不支持的文件格式: {file_ext}")
        
    except Exception as e:
        print(f"读取文件失败: {e}")
        import traceback
        traceback.print_exc()
    
    print(f"从文件 {os.path.basename(excel_file)} 提取到 {len(all_records)} 条记录")
    return all_records

def clean_data(records):
    """清洗数据"""
    cleaned_records = []
    
    print(f"清洗前记录数量: {len(records)}")
    
    for record in records:
        # 打印记录信息用于调试
        print(f"处理记录: 客户={record['client_name']}, 状态={record['apply_statue']}, 日期={record['apply_date']}")
        
        # 清洗client_name：去除末尾的ASCII大写字母后缀（如份额类别A/B/C等）
        record['client_name'] = re.sub(r'[A-Z]+$', '', record['client_name']).strip()

        # 删除client_name是上海睿量私募基金管理有限公司的记录
        if record['client_name'] == '上海睿量私募基金管理有限公司':
            continue

        # 只保留apply_statue为'成功'或'确认成功'的记录（无状态列时不过滤）
        if record['apply_statue']:
            if record['apply_statue'] not in ['成功', '确认成功', '已确认', '确认', 'SUCCESS', 'success']:
                # 检查是否是日期格式，如果是，说明列名匹配错误
                date_str = str(record['apply_statue']).strip()
                if len(date_str) == 8 and date_str.isdigit():
                    print(f"  警告: 状态字段是日期格式，可能列名匹配错误: {date_str}")
                continue
        
        cleaned_records.append(record)
    
    print(f"清洗后记录数量: {len(cleaned_records)}")
    return cleaned_records

def organize_data(records):
    """整理数据为字典"""
    data_dict = {}
    
    for record in records:
        client_name = record['client_name']
        apply_date = record['apply_date']
        
        if client_name not in data_dict:
            data_dict[client_name] = []
        
        if apply_date not in data_dict[client_name]:
            data_dict[client_name].append(apply_date)
    
    # 对日期进行排序
    for client_name in data_dict:
        # 转换日期格式为datetime进行排序
        date_objects = []
        for date_str in data_dict[client_name]:
            try:
                date_obj = pd.to_datetime(date_str, format='%Y年%m月%d日')
                date_objects.append((date_obj, date_str))
            except Exception:
                pass
        
        # 按日期排序并返回原始格式
        sorted_dates = [date_str for _, date_str in sorted(date_objects)]
        data_dict[client_name] = sorted_dates
    
    return data_dict

def create_word_document(client_name, apply_dates, client_code):
    """创建Word文档"""
    doc = Document()
    
    # 第一行：关于client_name（居中，四号，仿宋）
    para1 = doc.add_paragraph()
    run1 = para1.add_run(f'关于{client_name}')
    run1.font.name = '仿宋'
    run1.font.size = Pt(14)  # 四号
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')  # 设置中文字体
    para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 第二行：触发关联交易的公告（居中，四号，仿宋）
    para2 = doc.add_paragraph()
    run2 = para2.add_run('触发关联交易的公告')
    run2.font.name = '仿宋'
    run2.font.size = Pt(14)  # 四号
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')  # 设置中文字体
    para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 第三行：尊敬的投资者：（靠左，小四，仿宋）
    para3 = doc.add_paragraph()
    run3 = para3.add_run('尊敬的投资者：')
    run3.font.name = '仿宋'
    run3.font.size = Pt(12)  # 小四
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')  # 设置中文字体
    para3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 第四行：正文内容（靠左，段前留两格，小四，仿宋）
    para4 = doc.add_paragraph()
    # 添加首行缩进（两格，约28磅）
    para4.paragraph_format.first_line_indent = Pt(28)
    dates_str = '、'.join(apply_dates)
    content = f'上海睿量私募基金管理有限公司（以下简称"我公司"）管理的{client_name}（基金备案编码：{client_code}）（以下简称"本基金"）于{dates_str}发生关联交易，根据《基金合同》信息披露内容和《私募投资基金信息披露管理办法》对重大事项披露事项之"发生重大关联交易"规定，我公司对本次关联交易进行披露，我司已建立健全关联交易管理制度，本笔关联交易已履行基金合同及公司内控约定的决策程序。我司承诺不存在以私募基金财产与关联方进行不正当交易或者利益输送，违反法律法规、基金合同的有关约定。'
    run4 = para4.add_run(content)
    run4.font.name = '仿宋'
    run4.font.size = Pt(12)  # 小四
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')  # 设置中文字体
    para4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 第五行：特此说明。（靠左，段前留两格，小四，仿宋）
    para5 = doc.add_paragraph()
    # 添加首行缩进（两格，约28磅）
    para5.paragraph_format.first_line_indent = Pt(28)
    run5 = para5.add_run('特此说明。')
    run5.font.name = '仿宋'
    run5.font.size = Pt(12)  # 小四
    run5._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')  # 设置中文字体
    para5.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 添加空行
    doc.add_paragraph()
    
    # 第六行：上海睿量私募基金管理有限公司（靠右，小四，仿宋）
    para6 = doc.add_paragraph()
    run6 = para6.add_run('上海睿量私募基金管理有限公司')
    run6.font.name = '仿宋'
    run6.font.size = Pt(12)  # 小四
    run6._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')  # 设置中文字体
    para6.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # 第七行：最新apply_date（靠右，小四，仿宋）
    para7 = doc.add_paragraph()
    latest_date = apply_dates[-1] if apply_dates else ''
    run7 = para7.add_run(latest_date)
    run7.font.name = '仿宋'
    run7.font.size = Pt(12)  # 小四
    run7._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')  # 设置中文字体
    para7.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    return doc

def process_excel_files(base_dir=None):
    """
    处理所有Excel文件并生成关联交易公告函Word文档。

    Args:
        base_dir: Excel文件所在的根目录。如果为None，则使用脚本所在目录（本地运行）。
    """
    try:
        if base_dir is None:
            base_dir = os.path.dirname(os.path.abspath(__file__))

        # 切换到工作目录
        os.chdir(base_dir)

        # 创建关联交易文件夹
        output_dir = os.path.join(base_dir, '关联交易')
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        print("开始查找并处理所有Excel文件...")
        print("=" * 50)

        # 查找所有Excel文件
        excel_files = find_all_excel_files(base_dir)
        print(f"共找到 {len(excel_files)} 个Excel文件")

        if not excel_files:
            print("未找到任何Excel文件，请检查上传的文件。")
            return {"success": False, "message": "未找到Excel文件", "count": 0}

        all_records = []

        for excel_file in excel_files:
            try:
                print(f"处理文件: {os.path.basename(excel_file)}")
                records = extract_info(excel_file)
                all_records.extend(records)
            except Exception as e:
                print(f"处理文件 {excel_file} 失败: {e}")
                continue

        # 数据清洗
        cleaned_records = clean_data(all_records)
        print(f"清洗后共得到 {len(cleaned_records)} 条有效记录")

        # 整理数据
        data_dict = organize_data(cleaned_records)
        print(f"共整理出 {len(data_dict)} 个基金的关联交易记录")

        # 生成Word文档
        print(f"开始生成Word文档，共 {len(data_dict)} 个基金")
        generated_count = 0
        missing_code_count = 0
        missing_code_names = []

        for client_name, apply_dates in data_dict.items():
            print(f"处理基金: {client_name}")
            print(f"  交易日期: {apply_dates}")

            if client_name in client_code_map:
                client_code = client_code_map[client_name]
                print(f"  基金备案编码: {client_code}")
                try:
                    doc = create_word_document(client_name, apply_dates, client_code)
                    # 生成文件名
                    safe_client_name = "".join(c for c in client_name if c.isalnum() or c in (' ', '-', '_')).rstrip()
                    # 获取最新日期
                    latest_date = apply_dates[-1] if apply_dates else ''
                    # 提取日期部分（去除年、月、日等汉字）
                    date_str = "".join(c for c in latest_date if c.isdigit())
                    doc_name = f"{safe_client_name}关联交易公告函-{date_str}.docx"
                    doc_path = os.path.join(output_dir, doc_name)
                    doc.save(doc_path)
                    print(f"  生成文档: {doc_name}")
                    generated_count += 1
                except Exception as e:
                    print(f"  生成文档失败: {e}")
                    import traceback
                    traceback.print_exc()
                    continue
            else:
                print(f"  未找到 {client_name} 的基金备案编码")
                missing_code_count += 1
                missing_code_names.append(client_name)

        print(f"\n生成文档统计:")
        print(f"  总基金数: {len(data_dict)}")
        print(f"  成功生成: {generated_count}")
        print(f"  缺少编码: {missing_code_count}")
        if missing_code_names:
            print(f"  缺少编码的产品: {', '.join(missing_code_names)}")

        print("\n" + "=" * 50)
        print("处理完成！")
        print(f"Word文档保存在: {output_dir}")

        return {
            "success": True,
            "message": f"成功生成 {generated_count} 个文档",
            "count": generated_count,
            "total_funds": len(data_dict),
            "missing_codes": missing_code_count,
            "missing_code_names": missing_code_names
        }

    except Exception as e:
        print(f"程序执行错误: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "message": str(e), "count": 0}

def main():
    """本地运行入口"""
    return process_excel_files()

if __name__ == "__main__":
    main()
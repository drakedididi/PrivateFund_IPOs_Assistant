import os
import pandas as pd
import random
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
import openpyxl
import xlrd
from datetime import datetime, timedelta

def find_all_excel_files(root_dir='.'):
    """递归查找所有文件夹中的Excel文件，返回(文件路径, 文件夹名称)的元组列表"""
    excel_files = []
    
    for root, _, files in os.walk(root_dir):
        for file in files:
            if file.lower().endswith(('.xlsx', '.xls')):
                full_path = os.path.join(root, file)
                # 获取文件夹名称（相对于root_dir的最后一级目录）
                folder_name = os.path.basename(root)
                excel_files.append((full_path, folder_name))
    
    if not excel_files:
        raise FileNotFoundError(f"在目录 '{root_dir}' 及其子目录中未找到Excel文件")
    
    return excel_files

def read_excel_file(excel_file):
    """根据文件类型选择合适的库读取Excel文件，正确处理日期格式"""
    _, ext = os.path.splitext(excel_file)
    
    if ext.lower() == '.xlsx':
        return read_xlsx_file(excel_file)
    elif ext.lower() == '.xls':
        return read_xls_file(excel_file)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")

def read_xlsx_file(excel_file):
    """使用openpyxl读取.xlsx文件"""
    wb = openpyxl.load_workbook(excel_file, data_only=True)
    ws = wb.active
    
    # 查找表头行
    header_row = None
    for row in range(1, 8):  # 尝试前7行
        row_values = [cell.value for cell in ws[row]]
        if any('客户名称' in str(val) or '投资者名称' in str(val) for val in row_values if val):
            header_row = row
            break
    
    if header_row is None:
        raise ValueError(f"无法在Excel文件中找到表头: {excel_file}")
    
    # 构建列名映射
    column_mapping = {
        'client_name': ['客户名称', '投资者名称'],
        'product_name': ['产品名称', '基金名称'],
        'apply_date': ['申请日期'],
        'business_type': ['业务类型'],
        'confirm_share': ['确认份额', '申请份额'],
        'confirm_amount': ['确定金额', '确认金额'],
        'apply_amount': ['申请申购金额', '申请金额']
    }
    
    found_columns = {}
    for col_idx, cell in enumerate(ws[header_row]):
        if cell.value:
            cell_value = str(cell.value).strip()
            for target_col, possible_names in column_mapping.items():
                if cell_value in possible_names:
                    found_columns[target_col] = col_idx
                    break
    
    # 验证所有必要的列都找到
    required_columns = ['client_name', 'product_name', 'apply_date', 'business_type']
    missing = [k for k in required_columns if k not in found_columns]
    if missing:
        raise ValueError(f"无法找到必要的列: {missing}。文件路径: {excel_file}")
    
    # 提取数据
    all_records = []
    for row in range(header_row + 1, ws.max_row + 1):
        # 基本信息
        client_name = ws.cell(row=row, column=found_columns['client_name'] + 1).value
        product_name = ws.cell(row=row, column=found_columns['product_name'] + 1).value
        business_type = ws.cell(row=row, column=found_columns['business_type'] + 1).value
        
        # 跳过空行或无效行
        if not client_name or not product_name or not business_type:
            continue
        
        # 数据清洗
        result = {
            'client_name': str(client_name).strip(),
            'product_name': str(product_name).strip(),
            'business_type': str(business_type).strip()
        }
        
        # 处理申请日期 - 使用openpyxl正确处理
        date_cell = ws.cell(row=row, column=found_columns['apply_date'] + 1)
        date_val = date_cell.value
        
        if date_val:
            try:
                # 处理Excel日期对象
                if isinstance(date_val, datetime):
                    result['apply_date'] = date_val.strftime('%Y%m%d')
                # 处理Excel序列号
                elif isinstance(date_val, (int, float)):
                    # Excel日期序列号从1900年1月1日开始
                    base_date = datetime(1899, 12, 30)
                    converted_date = base_date + timedelta(days=date_val)
                    # 检查日期是否合理（大于2000年）
                    if converted_date.year > 2000:
                        result['apply_date'] = converted_date.strftime('%Y%m%d')
                    else:
                        # 尝试作为字符串处理
                        date_str = str(date_val).strip()
                        # 尝试yyyymmdd格式
                        if len(date_str) == 8 and date_str.isdigit():
                            result['apply_date'] = date_str
                        else:
                            result['apply_date'] = date_str
                # 处理字符串格式
                else:
                    date_str = str(date_val).strip().replace(' ', '').replace(',', '').replace('¥', '').replace('￥', '')
                    # 尝试yyyymmdd格式
                    if len(date_str) == 8 and date_str.isdigit():
                        result['apply_date'] = date_str
                    else:
                        # 尝试其他日期格式
                        date_formats = ['%Y-%m-%d', '%Y/%m/%d', '%d/%m/%Y', '%d-%m-%Y', '%Y年%m月%d日', '%Y.%m.%d']
                        for fmt in date_formats:
                            try:
                                result['apply_date'] = datetime.strptime(date_str, fmt).strftime('%Y%m%d')
                                break
                            except Exception:
                                continue
                        else:
                            result['apply_date'] = date_str
            except Exception:
                result['apply_date'] = str(date_val).strip()
        else:
            result['apply_date'] = None
        
        # 处理申请份额 - 支持多种格式
        if 'confirm_share' in found_columns:
            share_val = ws.cell(row=row, column=found_columns['confirm_share'] + 1).value
            if share_val:
                try:
                    # 尝试直接转换
                    result['confirm_share'] = float(share_val)
                except Exception:
                    # 处理字符串格式的份额
                    try:
                        # 去除非数字字符
                        share_str = str(share_val).strip().replace(',', '').replace(' ', '').replace('份', '').replace('元', '')
                        result['confirm_share'] = float(share_str)
                    except Exception:
                        result['confirm_share'] = share_val
            else:
                result['confirm_share'] = None
        else:
            result['confirm_share'] = None
        
        # 处理确定金额 - 支持多种格式
        if 'confirm_amount' in found_columns:
            amount_val = ws.cell(row=row, column=found_columns['confirm_amount'] + 1).value
            if amount_val:
                try:
                    # 尝试直接转换
                    result['confirm_amount'] = float(amount_val)
                except Exception:
                    # 处理字符串格式的金额
                    try:
                        # 去除非数字字符
                        amount_str = str(amount_val).strip().replace(',', '').replace(' ', '').replace('元', '').replace('¥', '').replace('￥', '')
                        result['confirm_amount'] = float(amount_str)
                    except Exception:
                        result['confirm_amount'] = amount_val
            else:
                result['confirm_amount'] = None
        else:
            result['confirm_amount'] = None
        
        # 处理申请金额 - 支持多种格式
        if 'apply_amount' in found_columns:
            apply_amount_val = ws.cell(row=row, column=found_columns['apply_amount'] + 1).value
            if apply_amount_val:
                try:
                    # 尝试直接转换
                    result['apply_amount'] = float(apply_amount_val)
                except Exception:
                    # 处理字符串格式的金额
                    try:
                        # 去除非数字字符
                        apply_amount_str = str(apply_amount_val).strip().replace(',', '').replace(' ', '').replace('元', '').replace('¥', '').replace('￥', '')
                        result['apply_amount'] = float(apply_amount_str)
                    except Exception:
                        result['apply_amount'] = apply_amount_val
            else:
                result['apply_amount'] = None
        else:
            result['apply_amount'] = None
        
        all_records.append(result)
    
    wb.close()
    
    if not all_records:
        raise ValueError(f"Excel文件中没有有效数据。文件路径: {excel_file}")
    
    return all_records

def read_xls_file(excel_file):
    """使用xlrd读取.xls文件"""
    wb = xlrd.open_workbook(excel_file)
    ws = wb.sheet_by_index(0)
    
    # 查找表头行
    header_row = None
    for row in range(min(7, ws.nrows)):  # 尝试前7行
        row_values = ws.row_values(row)
        if any('客户名称' in str(val) or '投资者名称' in str(val) for val in row_values if val):
            header_row = row
            break
    
    if header_row is None:
        raise ValueError(f"无法在Excel文件中找到表头: {excel_file}")
    
    # 构建列名映射
    column_mapping = {
        'client_name': ['客户名称', '投资者名称'],
        'product_name': ['产品名称', '基金名称'],
        'apply_date': ['申请日期'],
        'business_type': ['业务类型'],
        'confirm_share': ['确认份额', '申请份额'],
        'confirm_amount': ['确定金额', '确认金额'],
        'apply_amount': ['申请申购金额', '申请金额']
    }
    
    found_columns = {}
    header_values = ws.row_values(header_row)
    for col_idx, cell_value in enumerate(header_values):
        if cell_value:
            cell_value_str = str(cell_value).strip()
            for target_col, possible_names in column_mapping.items():
                if cell_value_str in possible_names:
                    found_columns[target_col] = col_idx
                    break
    
    # 验证所有必要的列都找到
    required_columns = ['client_name', 'product_name', 'apply_date', 'business_type']
    missing = [k for k in required_columns if k not in found_columns]
    if missing:
        raise ValueError(f"无法找到必要的列: {missing}。文件路径: {excel_file}")
    
    # 提取数据
    all_records = []
    for row in range(header_row + 1, ws.nrows):
        # 基本信息
        client_name = ws.cell_value(row, found_columns['client_name'])
        product_name = ws.cell_value(row, found_columns['product_name'])
        business_type = ws.cell_value(row, found_columns['business_type'])
        
        # 跳过空行或无效行
        if not client_name or not product_name or not business_type:
            continue
        
        # 数据清洗
        result = {
            'client_name': str(client_name).strip(),
            'product_name': str(product_name).strip(),
            'business_type': str(business_type).strip()
        }
        
        # 处理申请日期 - 使用xlrd正确处理
        date_val = ws.cell_value(row, found_columns['apply_date'])
        
        if date_val:
            try:
                # 处理Excel日期序列号
                if isinstance(date_val, (int, float)):
                    # 检查是否为日期序列号
                    if 60 <= date_val <= 60000:  # 合理的日期范围
                        # Excel日期序列号从1900年1月1日开始
                        base_date = datetime(1899, 12, 30)
                        converted_date = base_date + timedelta(days=date_val)
                        # 检查日期是否合理（大于2000年）
                        if converted_date.year > 2000:
                            result['apply_date'] = converted_date.strftime('%Y%m%d')
                        else:
                            # 尝试作为字符串处理
                            date_str = str(date_val).strip()
                            # 尝试yyyymmdd格式
                            if len(date_str) == 8 and date_str.isdigit():
                                result['apply_date'] = date_str
                            else:
                                result['apply_date'] = date_str
                    else:
                        # 尝试作为字符串处理
                        date_str = str(date_val).strip()
                        # 尝试yyyymmdd格式
                        if len(date_str) == 8 and date_str.isdigit():
                            result['apply_date'] = date_str
                        else:
                            # 尝试其他日期格式
                            date_formats = ['%Y-%m-%d', '%Y/%m/%d', '%d/%m/%Y', '%d-%m-%Y', '%Y年%m月%d日', '%Y.%m.%d']
                            for fmt in date_formats:
                                try:
                                    result['apply_date'] = datetime.strptime(date_str, fmt).strftime('%Y%m%d')
                                    break
                                except Exception:
                                    continue
                            else:
                                result['apply_date'] = date_str
                # 处理字符串格式
                else:
                    date_str = str(date_val).strip().replace(' ', '').replace(',', '').replace('¥', '').replace('￥', '')
                    # 尝试yyyymmdd格式
                    if len(date_str) == 8 and date_str.isdigit():
                        result['apply_date'] = date_str
                    else:
                        # 尝试其他日期格式
                        date_formats = ['%Y-%m-%d', '%Y/%m/%d', '%d/%m/%Y', '%d-%m-%Y', '%Y年%m月%d日', '%Y.%m.%d']
                        for fmt in date_formats:
                            try:
                                result['apply_date'] = datetime.strptime(date_str, fmt).strftime('%Y%m%d')
                                break
                            except Exception:
                                continue
                        else:
                            result['apply_date'] = date_str
            except Exception:
                result['apply_date'] = str(date_val).strip()
        else:
            result['apply_date'] = None
        
        # 处理申请份额 - 支持多种格式
        if 'confirm_share' in found_columns:
            share_val = ws.cell_value(row, found_columns['confirm_share'])
            if share_val:
                try:
                    # 尝试直接转换
                    result['confirm_share'] = float(share_val)
                except Exception:
                    # 处理字符串格式的份额
                    try:
                        # 去除非数字字符
                        share_str = str(share_val).strip().replace(',', '').replace(' ', '').replace('份', '').replace('元', '')
                        result['confirm_share'] = float(share_str)
                    except Exception:
                        result['confirm_share'] = share_val
            else:
                result['confirm_share'] = None
        else:
            result['confirm_share'] = None
        
        # 处理确定金额 - 支持多种格式
        if 'confirm_amount' in found_columns:
            amount_val = ws.cell_value(row, found_columns['confirm_amount'])
            if amount_val:
                try:
                    # 尝试直接转换
                    result['confirm_amount'] = float(amount_val)
                except Exception:
                    # 处理字符串格式的金额
                    try:
                        # 去除非数字字符
                        amount_str = str(amount_val).strip().replace(',', '').replace(' ', '').replace('元', '').replace('¥', '').replace('￥', '')
                        result['confirm_amount'] = float(amount_str)
                    except Exception:
                        result['confirm_amount'] = amount_val
            else:
                result['confirm_amount'] = None
        else:
            result['confirm_amount'] = None
        
        # 处理申请金额 - 支持多种格式
        if 'apply_amount' in found_columns:
            apply_amount_val = ws.cell_value(row, found_columns['apply_amount'])
            if apply_amount_val:
                try:
                    # 尝试直接转换
                    result['apply_amount'] = float(apply_amount_val)
                except Exception:
                    # 处理字符串格式的金额
                    try:
                        # 去除非数字字符
                        apply_amount_str = str(apply_amount_val).strip().replace(',', '').replace(' ', '').replace('元', '').replace('¥', '').replace('￥', '')
                        result['apply_amount'] = float(apply_amount_str)
                    except Exception:
                        result['apply_amount'] = apply_amount_val
            else:
                result['apply_amount'] = None
        else:
            result['apply_amount'] = None
        
        all_records.append(result)
    
    wb.release_resources()
    
    if not all_records:
        raise ValueError(f"Excel文件中没有有效数据。文件路径: {excel_file}")
    
    return all_records

def extract_info(excel_file):
    """从Excel文件中提取信息，根据文件类型选择合适的库处理日期格式"""
    return read_excel_file(excel_file)

def get_effective_amount(record):
    """获取有效金额，优先确认金额，无确认金额时回退到申请金额"""
    amount = record.get('confirm_amount')
    if amount is None:
        amount = record.get('apply_amount')
    return amount


def clean_data(records):
    """数据清洗，过滤掉不符合条件的记录"""
    cleaned_records = []
    for record in records:
        # 过滤掉客户名称为"总计"的记录
        if record['client_name'] == '总计':
            continue

        # 过滤掉金额小于1000000（100万）或无金额的记录
        # 优先使用确认金额，无确认金额时回退到申请金额
        effective_amount = get_effective_amount(record)
        amount_ok = False
        if isinstance(effective_amount, (int, float)):
            if effective_amount >= 1000000:
                amount_ok = True
        elif effective_amount is not None:
            # 非数值类型（如字符串）保留，维持原有行为
            amount_ok = True

        # 对于赎回类业务，申请金额可能为0，改用确认份额判断
        if not amount_ok and '赎回' in record.get('business_type', ''):
            share = record.get('confirm_share')
            if isinstance(share, (int, float)) and share >= 1000000:
                amount_ok = True

        if not amount_ok:
            continue

        # 过滤掉申请日期为空的记录
        if not record['apply_date']:
            continue
        cleaned_records.append(record)
    return cleaned_records

def process_normal_business(records, folder_name, output_dir=None):
    """处理普通业务类（申购、赎回）"""
    processed_count = 0
    for record in records:
        try:
            doc_name = create_word_document(record, folder_name, output_dir)
            processed_count += 1
        except Exception as e:
            continue
    return processed_count

def process_special_business(conversion_out_records, conversion_in_records, folder_name, output_dir=None):
    """处理特殊业务类（基金转换）"""
    processed_count = 0
    paired_indices = set()
    
    for i, out_record in enumerate(conversion_out_records):
        if i in paired_indices:
            continue
        
        for j, in_record in enumerate(conversion_in_records):
            if j in paired_indices:
                continue
            
            # 配对条件：客户名称相同且确定金额相近
            client_name_match = out_record['client_name'] == in_record['client_name']
            
            # 处理金额比较，确保类型一致（优先确认金额，回退到申请金额）
            out_amount = get_effective_amount(out_record)
            in_amount = get_effective_amount(in_record)
            
            # 尝试将金额转换为数字进行比较
            try:
                out_amount_num = float(out_amount) if out_amount is not None else 0
                in_amount_num = float(in_amount) if in_amount is not None else 0
                amount_match = abs(out_amount_num - in_amount_num) < 0.01  # 允许小误差
            except Exception:
                # 如果转换失败，使用字符串比较
                amount_match = str(out_amount) == str(in_amount)
            
            # 增加日期匹配条件
            date_match = out_record.get('apply_date') == in_record.get('apply_date')
            

            
            if client_name_match and amount_match and date_match:
                
                # 创建配对后的记录
                conversion_record = {
                    'client_name': out_record['client_name'],
                    'apply_date': out_record['apply_date'],
                    'business_type': '基金转换',
                    'confirm_share': out_record.get('confirm_share', out_record.get('apply_share')),
                    'confirm_amount': out_record['confirm_amount'],
                    'product_name_out': out_record['product_name'],
                    'product_name_in': in_record['product_name'],
                    'is_conversion': True
                }
                
                try:
                    doc_name = create_word_document(conversion_record, folder_name, output_dir)
                    processed_count += 1
                    
                    # 标记为已配对
                    paired_indices.add(i)
                    paired_indices.add(j)
                    break
                    
                except Exception as e:
                    continue
    
    return processed_count

def create_word_document(info, folder_name, output_dir=None):
    """创建Word文档，保存在脚本所在目录"""
    doc = Document()
    
    # 创建表格，一列五行
    table = doc.add_table(rows=5, cols=1)
    table.autofit = False
    table.columns[0].width = 5000000  # 设置列宽
    
    # 添加表格全部框线
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    
    for row in table.rows:
        for cell in row.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            for side in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{side}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '4')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), '000000')
                tcPr.append(border)
    
    # 第一行：关联交易决策机制（居中排列，加粗）
    cell1 = table.cell(0, 0)
    cell1.paragraphs[0].text = '关联交易决策机制'
    cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 设置垂直居中
    cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置1.5倍行距
    cell1.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run1 = cell1.paragraphs[0].runs[0]
    run1.font.name = '仿宋'
    run1.font.size = Pt(12)  # 小四
    run1.font.bold = True
    run1._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    # 第二行：决策日期（靠左，加粗）
    cell2 = table.cell(1, 0)
    cell2.paragraphs[0].text = f'决策日期：{info["apply_date"]}'
    cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    # 设置垂直居中
    cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置1.5倍行距
    cell2.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run2 = cell2.paragraphs[0].runs[0]
    run2.font.name = '仿宋'
    run2.font.size = Pt(12)  # 小四
    run2.font.bold = True
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    # 第三行：决策内容（靠左，加粗）
    cell3 = table.cell(2, 0)
    cell3.paragraphs[0].text = '决策内容：'
    cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    content = '决策内容：\n'
    content += f'1. 决策标的（产品）：{info["client_name"]}\n\n'
    
    # 处理基金转换的特殊情况
    if info.get('is_conversion', False):
        content += f'2. 拟投资标的（产品出）：{info["product_name_out"]}\n\n'
        content += f'   拟投资标的（产品入）：{info["product_name_in"]}\n\n'
    else:
        content += f'2. 拟投资标的（产品）：{info["product_name"]}\n\n'
    
    content += f'3. 拟投资日期：{info["apply_date"]}\n\n'
    
    # 处理拟投资定价
    if info['business_type'] == '申购' or info['business_type'] == '申购确认':
        content += f'4. 拟投资定价：申购 {get_effective_amount(info)} 元\n\n'
    elif info['business_type'] == '赎回' or info['business_type'] == '赎回确认':
        content += f'4. 拟投资定价：赎回 {info["confirm_share"]} 份\n\n'
    elif info['business_type'] == '金额赎回':
        content += f'4. 拟投资定价：赎回 {get_effective_amount(info)} 元\n\n'
    elif '转换' in info['business_type']:
        content += f'4. 拟投资定价：基金转换 {info["confirm_share"]} 份\n\n'
    
    # 添加决策依据
    content += '5. 决策依据：'
    
    # 根据客户名称选择不同的决策依据
    if info['client_name'] == '上海睿量私募基金管理有限公司':
        reasons = [
            '通过自营投资，公司能够对内部资产进行更灵活的配置，以适应市场变化和投资目标。',
            '公司可能根据特定的投资策略，如价值投资、成长投资或市场时机把握，进行自营投资以实现策略目标。',
            '将内部研发的交易模型、算法或策略通过自营投资进行实际应用，以验证和完善这些策略的有效性。',
            '利用自营投资作为工具，对冲投资组合中的市场风险、信用风险或其他相关风险。',
            '抓住市场中的短期交易机会或套利机会，通过自营投资获取额外收益。',
            '通过自营投资活动，公司可以寻求超越市场平均水平的收益，增强整体财务表现。',
        ]
    else:
        reasons = [
            '根据市场分析和预期，调整不同产品间的资产配置，以实现最优的投资组合平衡。',
            '确保跨产品的投资策略执行一致性，以维持整体投资风格的连贯性和策略的有效性。',
            '通过产品间的投资交易，分散特定资产或市场的风险，增强整体投资组合的稳健性。',
            '根据各产品的流动性需求，进行资金调配，确保各产品能够应对赎回压力或其他资金需求。',
            '通过内部产品间的交易减少交易成本，提高资本运用效率，增加整体收益。',
            '利用市场波动和时机，通过产品间的交易快速调整持仓，捕捉投资机会。',
            '确保跨产品交易符合监管要求，包括但不限于信息披露、交易规则和合规性标准。',
            '利用跨产品交易测试新的投资理念或策略，推动投资研究和产品创新。',
            '通过产品间的交易确保投资组合跟踪业绩基准，实现业绩的一致性。',
            '根据客户需求和偏好，通过产品间交易提供定制化服务或满足特定投资目标。',
            '通过产品间的交易提高资本使用效率，确保资本在不同投资机会中得到有效利用。',
            '定期对投资组合进行再平衡，以维持既定的风险收益特征和投资目标。',
            '执行跨产品的风险管理策略，包括对冲和风险转移，以保护投资组合免受不利市场变动的影响。',
            '根据宏观经济指标和趋势，通过产品间的交易调整投资组合，以适应经济环境的变化。',
            '考虑ESG因素，通过产品间的交易支持可持续发展和社会责任投资。',
            '投资于技术进步和创新领域，通过产品间的交易实现对新兴技术和行业的投资。',
            '针对不同产品特定投资目标，通过跨产品交易实现定制化投资策略。'
        ]
    
    content += random.choice(reasons)
    
    cell3.paragraphs[0].text = content
    cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    # 设置垂直居中
    cell3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置1.5倍行距
    cell3.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run3 = cell3.paragraphs[0].runs[0]
    run3.font.name = '仿宋'
    run3.font.size = Pt(12)  # 小四
    run3.font.bold = True
    run3._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    
    # 第四行：合规风控负责人签字（靠左，不加粗）
    cell4 = table.cell(3, 0)
    cell4.paragraphs[0].text = '合规风控负责人签字：\n'
    cell4.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell4.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置1.5倍行距
    cell4.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run4 = cell4.paragraphs[0].runs[0]
    run4.font.name = '仿宋'
    run4.font.size = Pt(12)  # 小四
    run4._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    
    # 第五行：执行董事签字（靠左，不加粗）
    cell5 = table.cell(4, 0)
    cell5.paragraphs[0].text = '执行董事签字：\n'
    cell5.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    cell5.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # 设置1.5倍行距
    cell5.paragraphs[0].paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run5 = cell5.paragraphs[0].runs[0]
    run5.font.name = '仿宋'
    run5.font.size = Pt(12)  # 小四
    run5._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
    
    # 生成文件名
    if info.get('is_conversion', False):
        # 基金转换的文件名
        safe_client = "".join(c for c in info['client_name'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_product_out = "".join(c for c in info['product_name_out'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_product_in = "".join(c for c in info['product_name_in'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
        doc_name = f"【{folder_name}】关联交易决策留档-{safe_product_out}转换{safe_product_in}-{safe_client}-{info.get('confirm_share', info.get('apply_share'))}份.docx"
    else:
        # 普通交易的文件名
        safe_product = "".join(c for c in info['product_name'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
        safe_client = "".join(c for c in info['client_name'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
        if info['business_type'] in ['申购', '申购确认']:
            doc_name = f"【{folder_name}】关联交易决策留档-{safe_product}-{safe_client}-申购{get_effective_amount(info)}元.docx"
        elif info['business_type'] in ['赎回', '赎回确认']:
            doc_name = f"【{folder_name}】关联交易决策留档-{safe_product}-{safe_client}-赎回{info.get('confirm_share', info.get('apply_share'))}份.docx"
        elif info['business_type'] == '金额赎回':
            doc_name = f"【{folder_name}】关联交易决策留档-{safe_product}-{safe_client}-赎回{get_effective_amount(info)}元.docx"
        else:
            doc_name = f"【{folder_name}】关联交易决策留档-{safe_product}-{safe_client}-{info['business_type']}.docx"
    
    # 保存文档
    doc_path = os.path.join(output_dir, doc_name) if output_dir else doc_name
    doc.save(doc_path)
    return doc_path

def process_excel_files(base_dir=None, output_dir=None):
    """
    处理所有Excel文件并生成Word文档。

    Args:
        base_dir: Excel文件所在的根目录。如果为None，则使用脚本所在目录（本地运行）。
    """
    original_cwd = os.getcwd()
    try:
        if base_dir is None:
            base_dir = os.path.dirname(os.path.abspath(__file__))
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        # 切换到工作目录，确保生成的Word文档保存在该目录
        os.chdir(base_dir)

        print("开始查找并处理所有Excel文件...")
        print("=" * 50)

        # 查找所有Excel文件
        excel_files = find_all_excel_files(base_dir)
        print(f"共找到 {len(excel_files)} 个Excel文件")

        if not excel_files:
            print("未找到任何Excel文件，请检查上传的文件。")
            return {"success": False, "message": "未找到Excel文件", "count": 0}

        # 统计文件夹数量和名称
        folders = set()
        for _, folder_name in excel_files:
            folders.add(folder_name)
        print(f"共涉及 {len(folders)} 个文件夹")
        print(f"文件夹名称：{', '.join(folders)}")
        print("=" * 50)

        total_processed = 0
        total_normal_records = 0
        total_special_records = 0

        for excel_file, folder_name in excel_files:
            try:
                print(f"处理文件夹: {folder_name}")
                print(f"处理文件: {os.path.basename(excel_file)}")

                # 提取信息
                records = extract_info(excel_file)

                # 数据清洗
                cleaned_records = clean_data(records)

                if not cleaned_records:
                    print("无有效记录，跳过")
                    continue

                # 分类记录
                conversion_in_records = []
                conversion_out_records = []
                normal_records = []

                for record in cleaned_records:
                    biz = record['business_type']
                    # 识别各种转换类型的业务（基金转换、份额转换等）
                    if '转换' in biz:
                        if '入' in biz:
                            conversion_in_records.append(record)
                        elif '出' in biz:
                            conversion_out_records.append(record)
                        else:
                            # 无明确方向的转换，默认为转出
                            conversion_out_records.append(record)
                    else:
                        normal_records.append(record)

                # 打印基金转换记录数量
                print(f'基金转换(入)记录: {len(conversion_in_records)}')
                print(f'基金转换(出)记录: {len(conversion_out_records)}')

                # 处理普通业务
                normal_count = process_normal_business(normal_records, folder_name, output_dir)
                total_normal_records += normal_count

                # 处理特殊业务
                special_count = process_special_business(conversion_out_records, conversion_in_records, folder_name, output_dir)
                total_special_records += special_count

                total_processed += 1
                print(f"处理完成，生成 {normal_count + special_count} 个Word文档")
                print("-" * 30)

            except Exception as e:
                print(f"处理失败: {e}")
                print("-" * 30)
                continue

        print("\n" + "=" * 50)
        print("处理完成！")
        print(f"成功处理 {total_processed}/{len(excel_files)} 个Excel文件")
        print(f"共生成 {total_normal_records + total_special_records} 个Word文档")
        print(f"其中：普通业务 {total_normal_records} 个，特殊业务（基金转换） {total_special_records} 个")

        if total_processed > 0:
            print(f"\nWord文档保存在目录: {base_dir}")

        return {
            "success": True,
            "message": f"成功处理 {total_processed} 个文件，生成 {total_normal_records + total_special_records} 个文档",
            "count": total_normal_records + total_special_records,
            "normal": total_normal_records,
            "special": total_special_records
        }

    except Exception as e:
        print(f"程序执行错误: {e}")
        import traceback
        traceback.print_exc()
        return {"success": False, "message": str(e), "count": 0}
    finally:
        os.chdir(original_cwd)

def main():
    """本地运行入口"""
    return process_excel_files()

if __name__ == "__main__":
    main()

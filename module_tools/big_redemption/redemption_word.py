import os
import json
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


def load_special_dates():
    """从项目 data 目录读取节假日日期。"""
    project_root = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', '..'))
    holidays_path = os.path.join(project_root, 'data', 'trading_holidays.json')
    with open(holidays_path, 'r', encoding='utf-8') as f:
        return set(json.load(f).get('vacation_dates', []))


SPECIAL_DATES = load_special_dates()


# 修改现有的find_excel_file函数，或者添加新函数
def find_all_excel_files(root_dir='.'):
    """递归查找所有文件夹中的Excel文件"""
    excel_files = []

    for root, _, files in os.walk(root_dir):
        for file in files:
            if file.lower().endswith(('.xlsx', '.xls')):
                full_path = os.path.join(root, file)
                excel_files.append(full_path)

    if not excel_files:
        raise FileNotFoundError(f"在目录 '{root_dir}' 及其子目录中未找到Excel文件")

    return excel_files

def extract_info(excel_file):
    """从Excel文件中提取信息，支持列名在不同行、申请日期为空、确认日期为自定义格式的情况"""
    # 尝试更多的表头行数
    header_rows = [0, 1, 2, 3, 4, 5, 6]  # 尝试前5行作为表头
    df = None
    found_columns = None

    for header_row in header_rows:
        try:
            # 读取Excel文件，跳过指定行数作为表头
            df = pd.read_excel(excel_file, header=header_row)

            # 检查数据是否为空
            if df.empty:
                continue

            # 动态查找列名
            column_mapping = {
                'product_name': ['产品名称', 'product name', '产品名'],
                'product_code': ['ta代码', '产品代码', 'product code', 'ta code', '产品代码'],
                'date': ['确认日期', '日期', '申请日期', '开放日','date','业务申请日期','交易申请日期'],
                'redemption_ratio': ['巨额赎回确认比例', '赎回比例', '实际净赎回比例', '巨额赎回实际比例', '巨额赎回发生比例','redemption ratio','净赎回比例','实际赎回确认比例']
            }

            found_columns = {}
            date_type = None  # 记录日期类型：确认日期或申请日期
            for target_col, possible_names in column_mapping.items():
                for name in possible_names:
                    for col in df.columns:
                        col_str = str(col).strip()
                        col_lower = col_str.lower()
                        if col_str == name or name.lower() == col_lower:
                            found_columns[target_col] = col
                            # 记录日期类型
                            if target_col == 'date':
                                if col_str == '确认日期':
                                    date_type = '确认日期'
                                elif col_str == '申请日期':
                                    date_type = '申请日期'
                            break
                    if target_col in found_columns:
                        break

            # 验证所有必要的列都找到
            required_columns = ['product_name', 'product_code', 'date']
            missing = [k for k in required_columns if k not in found_columns]
            if not missing:
                break  # 找到所有必要列，退出循环

        except Exception as e:
            continue

    if not found_columns or df is None:
        raise ValueError(f"无法找到必要的列。文件路径: {excel_file}")

    # 处理所有行数据，返回列表
    all_records = []
    for i in range(len(df)):
        data = df.iloc[i]

        # 跳过空行或无效行
        product_name = str(data.get(found_columns.get('product_name', ''))).strip()
        product_code = str(data.get(found_columns.get('product_code', ''))).strip()
        if not product_name or product_name in ['nan', 'None'] or not product_code or product_code in ['nan', 'None']:
            continue

        # 数据清洗
        result = {
            'product_name': product_name,
            'product_code': product_code,
            'redemption_ratio': data.get(found_columns.get('redemption_ratio', ''), '')
        }

        # 处理日期 - 支持自定义格式
        date_val = data.get(found_columns.get('date', ''))

        # 处理YYYYMMDD格式的整数/浮点数（如20260518.0 → 2026-05-18）
        if isinstance(date_val, (int, float, np.integer, np.floating)) and not pd.isna(date_val):
            date_str = str(int(date_val))
            if len(date_str) == 8 and date_str.isdigit():
                date_val = f"{date_str[:4]}-{date_str[4:6]}-{date_str[6:8]}"

        if pd.isna(date_val):
            result['date'] = None
        else:
            try:
                # 尝试直接转换日期
                date_obj = pd.to_datetime(date_val)

                # 检查是否需要减1天
                if date_type == '确认日期' or date_obj.strftime('%Y-%m-%d') in SPECIAL_DATES:
                    date_obj = date_obj - pd.Timedelta(days=1)

                result['date'] = date_obj.strftime('%Y-%m-%d')
            except Exception:
                # 处理自定义格式的日期
                try:
                    # 尝试不同的日期格式
                    date_str = str(date_val).strip()
                    # 尝试常见的日期格式
                    date_obj = None
                    for fmt in ['%Y-%m-%d', '%Y/%m/%d', '%d/%m/%Y', '%d-%m-%Y', '%Y年%m月%d日']:
                        try:
                            date_obj = pd.to_datetime(date_str, format=fmt)
                            break
                        except Exception:
                            continue

                    if date_obj:
                        # 检查是否需要减1天
                        if date_type == '确认日期' or date_obj.strftime('%Y-%m-%d') in SPECIAL_DATES:
                            date_obj = date_obj - pd.Timedelta(days=1)

                        result['date'] = date_obj.strftime('%Y-%m-%d')
                    else:
                        # 如果所有格式都失败，使用原始值
                        result['date'] = date_str
                except Exception:
                    result['date'] = None

        all_records.append(result)

    if not all_records:
        raise ValueError(f"Excel文件中没有有效数据。文件路径: {excel_file}")

    return all_records

def create_word_document(info, output_dir=None):
    """创建Word文档，保存在脚本所在目录"""
    doc = Document()

    # 设置标题
    title = doc.add_paragraph(f'关于{info["product_name"]}')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.runs[0]
    run.font.name = '仿宋'
    run.font.size = Pt(14)  # 四号
    # 添加中文字体支持
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    subtitle = doc.add_paragraph('触发巨额赎回的公告')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.runs[0]
    run.font.name = '仿宋'
    run.font.size = Pt(14)  # 四号
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # 添加正文
    investor_paragraph = doc.add_paragraph()
    run = investor_paragraph.add_run('尊敬的投资者：    ')
    run.font.name = '仿宋'
    run.font.size = Pt(12)  # 小四
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # 处理日期显示格式
    if info['date']:
        date_str = pd.to_datetime(info['date']).strftime('%Y年%m月%d日')
    else:
        date_str = "未指定日期"

    # 处理赎回比例显示格式 - 统一保留两位小数
    redemption_ratio = info.get('redemption_ratio', '')

    if isinstance(redemption_ratio, (int, float)):
        # 如果是数字，直接格式化
        redemption_ratio_str = f"{redemption_ratio:.2f}%"
    else:
        # 如果是字符串，先处理再格式化
        ratio_str = str(redemption_ratio).strip()
        has_percent = '%' in ratio_str or '％' in ratio_str
        # 已带百分号的值本身就是百分比，只去除符号后保留两位小数
        if has_percent:
            ratio_str = ratio_str.replace('%', '').replace('％', '').strip()

        # 尝试转换为浮点数并保留两位小数
        try:
            ratio_float = float(ratio_str)
            redemption_ratio_str = f"{ratio_float:.2f}%"
        except ValueError:
            # 如果转换失败，使用原始值
            redemption_ratio_str = f"{ratio_str}%"

    content = f"上海睿量私募基金管理有限公司 （以下简称\"我公司\"）管理的{info['product_name']}（基金备案编码：{info['product_code']}）（以下简称\"本基金\"）于{date_str}发生巨额赎回，巨额赎回比例为{redemption_ratio_str}。根据《基金合同》信息披露内容和《私募投资基金信息披露管理办法》对重大事项披露事项之\"基金触发巨额赎回的\"规定，我公司已对本次巨额赎回按正常赎回程序办理全额赎回。"
    content_paragraph = doc.add_paragraph()
    run = content_paragraph.add_run('    ' + content)  # 段前空两格
    run.font.name = '仿宋'
    run.font.size = Pt(12)  # 小四
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    notice = doc.add_paragraph()
    run = notice.add_run('    特此说明。')  # 段前空两格
    run.font.name = '仿宋'
    run.font.size = Pt(12)  # 小四
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # 添加落款
    doc.add_paragraph()

    company = doc.add_paragraph('上海睿量私募基金管理有限公司')
    company.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = company.runs[0]
    run.font.name = '仿宋'
    run.font.size = Pt(12)  # 小四
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # 添加日期段落
    date_paragraph = doc.add_paragraph(date_str)
    date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_paragraph.runs[0]
    run.font.name = '仿宋'
    run.font.size = Pt(12)  # 小四
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # 生成文件名
    safe_product_name = "".join(c for c in info['product_name'] if c.isalnum() or c in (' ', '-', '_')).rstrip()
    # 添加strftime('%Y%m%d')形式的日期
    if info['date']:
        date_obj = pd.to_datetime(info['date'])
        date_suffix = date_obj.strftime('%Y%m%d')
    else:
        date_suffix = ''
    doc_name = f"{safe_product_name}-巨额赎回公告函-{date_suffix}.docx"

    # 保存文档（本地脚本默认保存在当前目录，Web API 可指定临时输出目录）
    if output_dir:
        os.makedirs(output_dir, exist_ok=True)
        doc_path = os.path.join(output_dir, doc_name)
    else:
        doc_path = doc_name
    doc.save(doc_path)
    return doc_path


def process_excel_files(input_dir, output_dir):
    """处理指定目录内的 Excel 文件，并把生成的 Word 文档写入 output_dir。"""
    os.makedirs(output_dir, exist_ok=True)
    excel_files = find_all_excel_files(input_dir)

    stats = {
        "excel_total": len(excel_files),
        "excel_processed": 0,
        "records_total": 0,
        "generated_total": 0,
        "generated_files": [],
        "errors": [],
    }

    for excel_file in excel_files:
        try:
            records = extract_info(excel_file)
            stats["excel_processed"] += 1
            stats["records_total"] += len(records)
        except Exception as exc:
            stats["errors"].append(f"{os.path.basename(excel_file)}: {exc}")
            continue

        for record in records:
            try:
                doc_path = create_word_document(record, output_dir=output_dir)
                stats["generated_files"].append(doc_path)
                stats["generated_total"] += 1
            except Exception as exc:
                product = record.get("product_name", "未知产品")
                stats["errors"].append(f"{product}: {exc}")

    if stats["generated_total"] == 0:
        details = "; ".join(stats["errors"][:5]) if stats["errors"] else "未生成任何 Word 文档"
        raise RuntimeError(details)

    return stats

def main():
    try:
        # 切换到脚本所在目录
        script_dir = os.path.dirname(os.path.abspath(__file__))
        os.chdir(script_dir)

        print("开始查找并处理所有Excel文件...")
        print("=" * 50)

        # 查找所有Excel文件
        excel_files = find_all_excel_files(script_dir)
        print(f"共找到 {len(excel_files)} 个Excel文件")

        total_processed = 0
        total_records = 0

        for excel_file in excel_files:
            print(f"\n处理文件: {os.path.basename(excel_file)}")
            print(f"完整路径: {excel_file}")
            print("-" * 30)

            try:
                # 提取信息（现在是多条记录）
                records = extract_info(excel_file)
                print(f"文件中包含 {len(records)} 条记录")

                # 为每条记录生成Word文档
                for i, record in enumerate(records):
                    try:
                        doc_name = create_word_document(record)
                        print(f"[OK] 记录 {i+1}: {record['product_name']} -> 已生成Word文档")
                        print(f"  文件名: {doc_name}")
                        total_records += 1

                    except Exception as e:
                        print(f"[FAIL] 记录 {i+1}: 生成Word文档失败 - {e}")
                        continue

                total_processed += 1

            except Exception as e:
                print(f"[FAIL] 处理文件失败: {e}")
                continue

        print("\n" + "=" * 50)
        print("处理完成！")
        print(f"成功处理 {total_processed}/{len(excel_files)} 个Excel文件")
        print(f"共生成 {total_records} 个Word文档")

        if total_processed > 0:
            print(f"\nWord文档保存在脚本所在目录: {script_dir}")

    except Exception as e:
        print(f"程序执行错误: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
